import os
import docx
from flask import Flask, request, redirect, url_for, send_from_directory, render_template, flash, jsonify
from werkzeug.utils import secure_filename
from deep_translator import GoogleTranslator
from celery import Celery

UPLOAD_FOLDER = 'uploads'
DOWNLOAD_FOLDER = 'downloads'
ALLOWED_EXTENSIONS = (['txt', 'docx'])

app = Flask(__name__)
app.config['SECRET_KEY'] = 'top-secret!'
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['DOWNLOAD_FOLDER'] = DOWNLOAD_FOLDER

# Celery configuration
app.config['CELERY_BROKER_URL'] = 'redis://localhost:6379/0'
app.config['CELERY_RESULT_BACKEND'] = 'redis://localhost:6379/0'

# Initialize Celery
celery = Celery(app.name, result_backend=app.config['CELERY_RESULT_BACKEND'], broker=app.config['CELERY_BROKER_URL'])


def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1] in ALLOWED_EXTENSIONS

info = []
@app.route('/', methods=['GET', 'POST'])
def index():
    global info
    if request.method == 'POST':
        if 'file' not in request.files:
            print('No file attached in request')
            return redirect(request.url)
        file = request.files['file']
        if file.filename == '':
            print('No file selected')
            return redirect(request.url)
        if file and allowed_file(file.filename):
            filename = secure_filename(file.filename)
            try:
                file.save(os.path.join(app.config['UPLOAD_FOLDER'], filename))
            except:
                print('sdfsdsdfsdf')
            flash('Текст из файла: {0} переводится!'.format(filename))
            task = translate.delay(app.config['UPLOAD_FOLDER'], filename)
            global info
            info.append(filename)
            info.append(task.id)
            return redirect(url_for('index', info=info))
    return render_template("index.html", info=info)


@celery.task
def translate(path, filename):
    if "docx" in filename:
        document = docx.Document(os.path.join(path, filename))
        new_document = docx.Document()
        list_text = []
        for paragraph in document.paragraphs:
            list_text.append(paragraph.text)
        text = '\n'.join(list_text)
        translated = GoogleTranslator(source='auto', target='en').translate(text)
        for string in translated.split("\n"):
            new_document.add_paragraph(string)
        styles = []
        alignments = []
        for paragraph in document.paragraphs:
            styles.append(paragraph.style.name)
            alignments.append(paragraph.paragraph_format.alignment)
        for i in range(len(new_document.paragraphs)):
            new_document.paragraphs[i].style = styles[i]
            new_document.paragraphs[i].paragraph_format.alignment = alignments[i]
        end_filename = "tr_" + filename
        new_document.save(os.path.join(app.config['DOWNLOAD_FOLDER'], end_filename))
    else:
        f = open(os.path.join(path, filename), 'r', encoding="utf-8")
        contents = f.read()
        translated = GoogleTranslator(source='auto', target='en').translate(contents)
        end_filename = "tr_" + filename
        with open(os.path.join(app.config['DOWNLOAD_FOLDER'], end_filename), 'w') as f:
            f.writelines(translated)
    return {'status': 'Перевод завершен!'}


@app.route('/uploads/<filename>')
def uploaded_file(filename):
    return send_from_directory(app.config['UPLOAD_FOLDER'], filename)


@app.route('/downloads/<filename>')
def downloaded_file(filename):
    return send_from_directory(app.config['DOWNLOAD_FOLDER'], filename, as_attachment=True)


@app.route('/status/<task_id>')
def taskstatus(task_id):
    task = translate.AsyncResult(task_id)
    if task.state == 'PENDING':
        response = {
            'state': task.state,
            'status': 'Переводится...'
        }
    elif task.state != 'FAILURE':
        response = {
            'state': task.state,
            'status': task.info.get('status', '')
        }
    else:
        response = {
            'state': task.state,
            'status': str(task.info),  # this is the exception raised
        }
    return jsonify(response)


if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5001)
